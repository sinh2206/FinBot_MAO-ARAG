from __future__ import annotations

import hashlib
import mimetypes
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from rag_engine.schema import Document


SUPPORTED_EXTENSIONS = {
    ".csv",
    ".doc",
    ".docx",
    ".htm",
    ".html",
    ".md",
    ".markdown",
    ".pdf",
    ".txt",
}


@dataclass(slots=True)
class LoadedFile:
    """Text content loaded from one source file."""

    path: Path
    text: str
    loader: str
    metadata: dict[str, object] = field(default_factory=dict)

    @property
    def id(self) -> str:
        source = str(self.metadata.get("relative_path") or self.path.name)
        digest = hashlib.sha1(source.encode("utf-8")).hexdigest()[:12]
        return f"doc-{digest}"

    def to_document(self) -> Document:
        metadata = {
            "source": str(self.path),
            "source_file": self.path.name,
            "file_name": self.path.name,
            "file_stem": self.path.stem,
            "extension": self.path.suffix.lower(),
            "loader": self.loader,
            "mime_type": mimetypes.guess_type(str(self.path))[0],
            "text_length": len(self.text),
            **self.metadata,
        }
        return Document(id=self.id, text=self.text, metadata=metadata)


def load_directory(
    directory: str | Path,
    *,
    recursive: bool = True,
    extensions: Iterable[str] | None = None,
    clean: bool = True,
) -> list[LoadedFile]:
    """Load all supported files from a directory into normalized text."""

    root = Path(directory)
    if not root.exists():
        raise FileNotFoundError(f"Directory does not exist: {root}")
    if not root.is_dir():
        raise NotADirectoryError(f"Expected a directory: {root}")

    allowed = normalize_extensions(extensions or SUPPORTED_EXTENSIONS)
    pattern = "**/*" if recursive else "*"
    loaded: list[LoadedFile] = []
    for path in sorted(root.glob(pattern)):
        if not path.is_file() or path.suffix.lower() not in allowed:
            continue
        item = load_file(path, clean=clean)
        item.metadata["relative_path"] = str(path.relative_to(root))
        loaded.append(item)
    return loaded


def load_file(path: str | Path, *, clean: bool = True) -> LoadedFile:
    """Load one supported file into text."""

    source = Path(path)
    if not source.exists():
        raise FileNotFoundError(source)
    if not source.is_file():
        raise ValueError(f"Expected a file: {source}")

    suffix = source.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file extension: {suffix}")

    text, loader = convert_source(source)
    if clean:
        text = clean_text(text)
    return LoadedFile(path=source, text=text, loader=loader)


def convert_source(source: Path) -> tuple[str, str]:
    """Convert a source file using the same converters as raw-data conversion."""

    suffix = source.suffix.lower()
    try:
        from scripts.convert_raw_data import convert_file

        return convert_file(
            source,
            ocr_mode="auto",
            ocr_lang="vie+eng",
            ocr_dpi=150,
            min_extracted_chars=200,
        )
    except ImportError:
        return fallback_convert_source(source, suffix)


def fallback_convert_source(source: Path, suffix: str) -> tuple[str, str]:
    if suffix in {".txt", ".md", ".markdown"}:
        return source.read_text(encoding="utf-8", errors="ignore"), suffix.lstrip(".")
    if suffix == ".csv":
        return source.read_text(encoding="utf-8-sig", errors="ignore"), "csv_text"
    if suffix in {".html", ".htm"}:
        raw = source.read_text(encoding="utf-8", errors="ignore")
        return re.sub(r"<[^>]+>", " ", raw), "html_fallback"
    if suffix == ".pdf":
        try:
            from PyPDF2 import PdfReader
        except ImportError as exc:
            raise RuntimeError("PyPDF2 is required for PDF loading") from exc
        pages = []
        reader = PdfReader(str(source))
        for index, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"[page {index}]\n{page_text}")
        return "\n\n".join(pages), "pdf_text"
    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise RuntimeError("python-docx is required for DOCX loading") from exc
        document = DocxDocument(str(source))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n\n".join(paragraphs), "docx"
    raise RuntimeError(f"No fallback converter for {suffix}")


def clean_text(text: str) -> str:
    value = text.replace("\x00", " ")
    try:
        from scripts.convert_raw_data import normalize_vietnamese_ocr_text

        value = normalize_vietnamese_ocr_text(value)
    except ImportError:
        pass
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"[ \t\f\v]+", " ", value)
    value = re.sub(r" *\n *", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def normalize_extensions(values: Iterable[str]) -> set[str]:
    normalized = set()
    for value in values:
        item = value.lower().strip()
        if not item:
            continue
        normalized.add(item if item.startswith(".") else f".{item}")
    return normalized
