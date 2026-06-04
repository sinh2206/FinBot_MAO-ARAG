from __future__ import annotations

import argparse
import csv
import json
import re
import subprocess
import sys
from io import BytesIO
from pathlib import Path
from typing import Iterable


PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))


DEFAULT_RAW_DIR = PROJECT_ROOT / "data" / "raw_data"
DEFAULT_PROCESSED_DIR = PROJECT_ROOT / "data" / "processed_data"
DEFAULT_METADATA_DIR = PROJECT_ROOT / "data" / "metadata"

SUPPORTED_EXTENSIONS = {
    ".csv",
    ".doc",
    ".docx",
    ".htm",
    ".html",
    ".md",
    ".markdown",
    ".pdf",
    ".txt",
}

OCR_PHRASE_CORRECTIONS: tuple[tuple[re.Pattern[str], str], ...] = (
    (re.compile(r"\bS[06o]\s+Giao\s+dich\b", flags=re.IGNORECASE), "Sở Giao dịch"),
    (re.compile(r"\bSo\s+Giao\s+dich\b", flags=re.IGNORECASE), "Sở Giao dịch"),
    (re.compile(r"\bS&\s+Giao\s+dịch\b", flags=re.IGNORECASE), "Sở Giao dịch"),
    (re.compile(r"\bGiao\s+dich\b", flags=re.IGNORECASE), "Giao dịch"),
    (
        re.compile(
            r"\bCh(?:ứng|ung|ing|img|tmg|irng|i?rng|ủng)\s+"
            r"kh(?:oan|oán|oàn|oản|oãn|oạn|oén|oen|ocin|odn|o[aáàảãạăâeéèẻẽẹd]n)\b",
            flags=re.IGNORECASE,
        ),
        "Chứng khoán",
    ),
    (re.compile(r"\bCHUNG\s+KHOAN\b", flags=re.IGNORECASE), "Chứng khoán"),
    (re.compile(r"\bChứng\s+kh(?:oan|oán|oàn|o[aáàảãạăâeéèẻẽẹd]n)\b", flags=re.IGNORECASE), "Chứng khoán"),
    (re.compile(r"\bThanh\s+ph(?:e|é|o|ó|d|đ)\b", flags=re.IGNORECASE), "Thành phố"),
    (re.compile(r"\bThành\s+ph(?:o|ó|d|đ|e|é)\b", flags=re.IGNORECASE), "Thành phố"),
    (re.compile(r"\bTP\.\s*H[eéèẻẽẹo]\s+Chi\s+Minh\b", flags=re.IGNORECASE), "TP. Hồ Chí Minh"),
    (re.compile(r"\bTP\s+H[O0]\s+CH[IÍÌỈĨỊ]\s+MINH\b", flags=re.IGNORECASE), "TP. Hồ Chí Minh"),
    (re.compile(r"\bThanh\s+ph6\s+H[O0]\s+Chi\s+Minh\b", flags=re.IGNORECASE), "Thành phố Hồ Chí Minh"),
    (
        re.compile(
            r"\bThành phố\s+H(?:o|6|e|é|è|ẻ|ẽ|ẹ|ồ|ố|ổ|ỗ|ộ|ò|ó|ỏ|õ|ọ)\s+Chi\s+Minh\b",
            flags=re.IGNORECASE,
        ),
        "Thành phố Hồ Chí Minh",
    ),
    (
        re.compile(
            r"\bThanh phố\s+H(?:o|6|e|é|è|ẻ|ẽ|ẹ|ồ|ố|ổ|ỗ|ộ|ò|ó|ỏ|õ|ọ)\s+Chi\s+Minh\b",
            flags=re.IGNORECASE,
        ),
        "Thành phố Hồ Chí Minh",
    ),
    (re.compile(r"\bThanh pho\s+H(?:o|6|e|é|è|ẻ|ẽ|ẹ)\s+Chi\s+Minh\b", flags=re.IGNORECASE), "Thành phố Hồ Chí Minh"),
    (re.compile(r"\bHOCHIMINH\s+STOCK\s+EXCHANGE\b", flags=re.IGNORECASE), "Ho Chi Minh Stock Exchange"),
    (re.compile(r"\bUy\s+ban\s+Chứng\s+khoán\s+Nhà\s+nước\b", flags=re.IGNORECASE), "Ủy ban Chứng khoán Nhà nước"),
    (re.compile(r"\bOy\s+ban\s+Chứng\s+khoán\b", flags=re.IGNORECASE), "Ủy ban Chứng khoán"),
    (re.compile(r"\bUY\s+BAN\s+Chứng\s+khoán\s+NHA\s+NUOC\b", flags=re.IGNORECASE), "Ủy ban Chứng khoán Nhà nước"),
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Convert raw PDF/DOC/DOCX/CSV/TXT documents from data/raw_data into "
            "normalized UTF-8 TXT files in data/processed_data."
        )
    )
    parser.add_argument(
        "--raw_dir",
        "--input_dir",
        dest="raw_dir",
        default=str(DEFAULT_RAW_DIR),
        help="Input folder containing original raw files.",
    )
    parser.add_argument(
        "--processed_dir",
        "--output_dir",
        dest="processed_dir",
        default=str(DEFAULT_PROCESSED_DIR),
        help="Output folder for normalized .txt files.",
    )
    parser.add_argument(
        "--metadata_dir",
        default=str(DEFAULT_METADATA_DIR),
        help="Output folder for conversion metadata JSON.",
    )
    parser.add_argument(
        "--recursive",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Scan raw_dir recursively. Use --no-recursive to scan only one level.",
    )
    parser.add_argument("--overwrite", action="store_true", help="Overwrite existing .txt outputs.")
    parser.add_argument(
        "--ocr",
        choices=["auto", "always", "never"],
        default="auto",
        help="PDF OCR mode. auto runs OCR only when normal PDF text extraction is too short.",
    )
    parser.add_argument("--ocr_lang", default="vie+eng", help="Tesseract languages, for example vie+eng.")
    parser.add_argument("--ocr_dpi", type=int, default=150, help="PDF render DPI before OCR.")
    parser.add_argument(
        "--min_extracted_chars",
        type=int,
        default=200,
        help="Minimum PDF text chars required before OCR is skipped in auto mode.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    raw_dir = Path(args.raw_dir)
    processed_dir = Path(args.processed_dir)
    metadata_dir = Path(args.metadata_dir)

    if not raw_dir.exists():
        raise FileNotFoundError(f"Raw data folder does not exist: {raw_dir}")

    processed_dir.mkdir(parents=True, exist_ok=True)
    metadata_dir.mkdir(parents=True, exist_ok=True)

    files = collect_files(raw_dir, recursive=args.recursive)
    summary_items: list[dict[str, object]] = []
    converted = 0
    skipped = 0
    failed = 0

    for source in files:
        relative = source.relative_to(raw_dir)
        output = (processed_dir / relative).with_suffix(".txt")
        output.parent.mkdir(parents=True, exist_ok=True)

        if output.exists() and not args.overwrite:
            skipped += 1
            summary_items.append(make_summary_item(source, output, "skipped", "exists", text_chars=output.stat().st_size))
            continue

        try:
            text, method = convert_file(
                source,
                ocr_mode=args.ocr,
                ocr_lang=args.ocr_lang,
                ocr_dpi=args.ocr_dpi,
                min_extracted_chars=args.min_extracted_chars,
            )
            text = clean_text(text)
            if not text:
                raise RuntimeError("Converted text is empty")

            output.write_text(text + "\n", encoding="utf-8")
            converted += 1
            summary_items.append(make_summary_item(source, output, "converted", method, text_chars=len(text)))
        except Exception as exc:  # noqa: BLE001 - batch conversion should continue.
            failed += 1
            summary_items.append(make_summary_item(source, output, "failed", source.suffix.lower(), error=str(exc)))

    summary = {
        "raw_dir": str(raw_dir),
        "processed_dir": str(processed_dir),
        "total_files": len(files),
        "converted": converted,
        "skipped": skipped,
        "failed": failed,
        "ocr_mode": args.ocr,
        "ocr_lang": args.ocr_lang,
        "ocr_dpi": args.ocr_dpi,
        "min_extracted_chars": args.min_extracted_chars,
        "files": summary_items,
    }
    summary_path = metadata_dir / "raw_data_conversion_summary.json"
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"Total files: {len(files)}")
    print(f"Converted: {converted}, skipped: {skipped}, failed: {failed}")
    print(f"Processed TXT folder: {processed_dir}")
    print(f"Summary: {summary_path}")
    if failed:
        raise SystemExit(1)


def collect_files(raw_dir: Path, *, recursive: bool = True) -> list[Path]:
    pattern = "**/*" if recursive else "*"
    return [
        path
        for path in sorted(raw_dir.glob(pattern))
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]


def convert_file(
    source: Path,
    *,
    ocr_mode: str,
    ocr_lang: str,
    ocr_dpi: int,
    min_extracted_chars: int,
) -> tuple[str, str]:
    suffix = source.suffix.lower()
    if suffix == ".pdf":
        return convert_pdf(source, ocr_mode, ocr_lang, ocr_dpi, min_extracted_chars)
    if suffix == ".docx":
        return convert_docx(source), "docx"
    if suffix == ".doc":
        return convert_doc(source), "doc"
    if suffix in {".txt", ".md", ".markdown"}:
        return source.read_text(encoding="utf-8", errors="ignore"), suffix.lstrip(".")
    if suffix in {".html", ".htm"}:
        return convert_html(source), "html"
    if suffix == ".csv":
        return convert_csv(source), "csv"
    raise ValueError(f"Unsupported extension: {suffix}")


def convert_pdf(
    source: Path,
    ocr_mode: str,
    ocr_lang: str,
    ocr_dpi: int,
    min_extracted_chars: int,
) -> tuple[str, str]:
    extracted = extract_pdf_text(source)
    extracted_chars = len(extracted.strip())
    should_ocr = ocr_mode == "always" or (ocr_mode == "auto" and extracted_chars < min_extracted_chars)
    if not should_ocr:
        return extracted, "pdf_text"

    ocr_text = ocr_pdf(source, lang=ocr_lang, dpi=ocr_dpi)
    if ocr_text.strip():
        method = "pdf_ocr" if not extracted.strip() else "pdf_text_plus_ocr"
        return "\n\n".join(part for part in [extracted, ocr_text] if part.strip()), method
    return extracted, "pdf_text"


def extract_pdf_text(source: Path) -> str:
    try:
        from PyPDF2 import PdfReader
    except ImportError as exc:
        raise RuntimeError("PyPDF2 is required for PDF text extraction") from exc

    reader = PdfReader(str(source))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if page_text.strip():
            pages.append(f"[page {index}]\n{page_text}")
    return "\n\n".join(pages)


def ocr_pdf(source: Path, *, lang: str, dpi: int) -> str:
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        raise RuntimeError("OCR requires pymupdf, pytesseract and pillow") from exc

    scale = dpi / 72
    matrix = fitz.Matrix(scale, scale)
    pages: list[str] = []
    with fitz.open(str(source)) as document:
        for index, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            image = Image.open(BytesIO(pixmap.tobytes("png")))
            page_text = pytesseract.image_to_string(image, lang=lang)
            if page_text.strip():
                pages.append(f"[page {index} OCR]\n{page_text}")
    return "\n\n".join(pages)


def convert_docx(source: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX conversion") from exc

    document = Document(str(source))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_rows: list[str] = []
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                table_rows.append(" | ".join(values))
    return "\n\n".join(paragraphs + table_rows)


def convert_doc(source: Path) -> str:
    try:
        completed = subprocess.run(["antiword", str(source)], check=True, capture_output=True)
    except FileNotFoundError as exc:
        raise RuntimeError("antiword is required for legacy .doc conversion") from exc
    except subprocess.CalledProcessError as exc:
        stderr = exc.stderr.decode("utf-8", errors="ignore").strip()
        raise RuntimeError(f"antiword failed: {stderr}") from exc
    return completed.stdout.decode("utf-8", errors="ignore")


def convert_html(source: Path) -> str:
    from html.parser import HTMLParser

    class HTMLTextExtractor(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self.parts: list[str] = []

        def handle_data(self, data: str) -> None:
            if data.strip():
                self.parts.append(data.strip())

    parser = HTMLTextExtractor()
    parser.feed(source.read_text(encoding="utf-8", errors="ignore"))
    return "\n".join(parser.parts)


def convert_csv(source: Path) -> str:
    lines: list[str] = []
    with source.open("r", encoding="utf-8-sig", errors="ignore", newline="") as handle:
        reader = csv.reader(handle)
        for row in reader:
            values = [cell.strip() for cell in row if cell.strip()]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines)


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = normalize_vietnamese_ocr_text(text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_vietnamese_ocr_text(text: str) -> str:
    normalized = text
    for pattern, replacement in OCR_PHRASE_CORRECTIONS:
        normalized = pattern.sub(replacement, normalized)
    return normalized


def make_summary_item(
    source: Path,
    output: Path,
    status: str,
    method: str,
    *,
    text_chars: int = 0,
    error: str | None = None,
) -> dict[str, object]:
    item: dict[str, object] = {
        "source": str(source),
        "output": str(output),
        "status": status,
        "method": method,
        "text_chars": text_chars,
    }
    if error:
        item["error"] = error
    return item


if __name__ == "__main__":
    main()
